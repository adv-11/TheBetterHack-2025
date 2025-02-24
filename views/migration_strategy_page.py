import streamlit as st
from gitingest import ingest  
import asyncio
import sys
from docx import Document
from docx.shared import Pt
from agents.reviewer_agent import reviewer_agent
from agents.fragmenter_agent import fragmentor_agent
from agents.migrator_agent import migrator_agent

# Fix for Windows subprocess issue
if sys.platform == "win32":
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

def save_string_as_word(text, filename):
    """Creates a Word document with the provided text and saves it to filename."""
    document = Document()
    document.add_heading("Migration Strategy", level=1)
    paragraphs = text.split("\n\n")
    for para in paragraphs:
        p = document.add_paragraph(para)
        for run in p.runs:
            run.font.size = Pt(11)
    document.save(filename)

st.title("GitHub Repository Ingestor")
st.write("Enter a GitHub repository URL to fetch its summary, tree, and content.")

# Initialize session state for repository data, human inputs, and migration strategy approval
if "repo_ingested" not in st.session_state:
    st.session_state.repo_ingested = False
if "summary" not in st.session_state:
    st.session_state.summary = None
if "tree" not in st.session_state:
    st.session_state.tree = None
if "content" not in st.session_state:
    st.session_state.content = None
if "migration_strategy" not in st.session_state:
    st.session_state.migration_strategy = None
if "migration_approved" not in st.session_state:
    st.session_state.migration_approved = None
if "human_expectations" not in st.session_state:
    st.session_state.human_expectations = None
if "human_guidelines" not in st.session_state:
    st.session_state.human_guidelines = None

repo_url = st.text_input("GitHub Repository URL", "https://github.com/neeti-kurulkar/feedback-system")

if st.button("Ingest Repository"):
    with st.spinner("Fetching repository data..."):
        try:
            # Call ingest() inside the button callback to delay cloning until the button is clicked.
            summary, tree, content = ingest(repo_url)
            st.session_state.summary = summary
            st.session_state.tree = tree
            st.session_state.content = content
            st.session_state.repo_ingested = True
            st.success("Repository data fetched successfully.")
        except Exception as e:
            st.error(f"Error fetching repository: {e}")

# Show human input fields once repository data is ingested
if st.session_state.repo_ingested:
    st.session_state.human_expectations = st.text_area("Please provide your expectations from the code migration:")
    st.session_state.human_guidelines = st.text_area("Please provide the guidelines that the Agent needs to follow while migrating:")

    if st.button("Generate Migration Strategy"):
        if not st.session_state.human_expectations or not st.session_state.human_guidelines:
            st.info("Please provide both expectations and guidelines to proceed.")
        else:
            try:
                migration_strategy = reviewer_agent(
                    st.session_state.tree,
                    st.session_state.content,
                    st.session_state.human_expectations,
                    st.session_state.human_guidelines
                )
                st.session_state.migration_strategy = migration_strategy
                word_filename = "migration_strategy.docx"
                save_string_as_word(st.session_state.migration_strategy, word_filename)
                st.success("Migration strategy generated successfully.")
                with open(word_filename, "rb") as file:
                    st.download_button(
                        label="Download Migration Strategy",
                        data=file,
                        file_name=word_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                st.subheader("Generated Migration Strategy")
                # st.text(migration_strategy)
            except Exception as e:
                st.error(f"Error generating migration strategy: {e}")

# Show approval logic after the strategy is generated
if st.session_state.migration_strategy:
    st.markdown("---")
    st.subheader("Approve Migration Strategy")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Approve Strategy"):
            st.session_state.migration_approved = True
            st.success("✅ Migration strategy approved. Proceeding with execution...")
    with col2:
        if st.button("Reject Strategy"):
            st.session_state.migration_approved = False
            st.error("❌ Migration strategy rejected.")
            feedback = st.text_area("Please provide feedback on what needs improvement:")
            if st.button("Submit Feedback"):
                st.info("🔄 Regenerating strategy based on feedback...")
                st.write(f"Feedback received: {feedback}")
                # Optionally: Use feedback to refine prompts and regenerate the strategy

# Run fragmentor and migrator agents only if previous steps succeeded
if st.session_state.repo_ingested and st.session_state.content and st.session_state.tree and st.session_state.migration_strategy:
    try:
        fragmented_data = fragmentor_agent(st.session_state.tree, st.session_state.content, st.session_state.migration_strategy)
        migrated_code = migrator_agent(st.session_state.tree, st.session_state.content, fragmented_data, st.session_state.human_expectations, st.session_state.human_guidelines)
        st.write("Migrated Code:")
        st.code(migrated_code)
    except Exception as e:
        st.error(f"Error during migration: {e}")
